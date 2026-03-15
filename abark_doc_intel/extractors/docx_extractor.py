"""
DOCX Extractor — extracts text, sections, and tables from .docx files.
Uses python-docx.
"""

import logging
import os
from ..models import DocumentResult, DocFormat, ExtractedSection, ExtractedTable

logger = logging.getLogger(__name__)


class DOCXExtractor:
    def extract(self, path: str) -> DocumentResult:
        logger.info(f"Extracting DOCX: {path}")
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = Document(path)
        raw_text = ""
        sections = []
        tables = []

        for para in doc.paragraphs:
            raw_text += para.text + "\n"
            if para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                except ValueError:
                    level = 1
                sections.append(ExtractedSection(
                    title=para.text.strip(),
                    content="",
                    page=None,
                    level=level,
                ))

        for tbl in doc.tables:
            if not tbl.rows:
                continue
            headers = [cell.text.strip() for cell in tbl.rows[0].cells]
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows[1:]]
            tables.append(ExtractedTable(headers=headers, rows=rows))

        logger.info(f"DOCX extraction complete: {len(raw_text)} chars, {len(tables)} tables")
        return DocumentResult(
            source=path,
            format=DocFormat.DOCX,
            raw_text=raw_text.strip(),
            sections=sections,
            tables=tables,
            metadata={"filename": os.path.basename(path)},
        )
